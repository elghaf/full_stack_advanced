from pathlib import Path
from typing import Tuple, List
import PyPDF2
from pdf2image import convert_from_path
from docx import Document
import os
import shutil

class FileProcessor:
    def __init__(self):
        self.upload_dir = Path("uploads")
        self.preview_dir = Path("previews")
        
        # Create directories if they don't exist
        self.upload_dir.mkdir(exist_ok=True)
        self.preview_dir.mkdir(exist_ok=True)

    async def process_pdf(self, file_path: Path) -> Tuple[int, List[str]]:
        """Process PDF files and generate previews."""
        try:
            # Get page count
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                page_count = len(pdf.pages)

            # Generate preview images
            images = convert_from_path(str(file_path))
            preview_urls = []

            # Save preview images
            file_preview_dir = self.preview_dir / file_path.stem
            file_preview_dir.mkdir(exist_ok=True)

            for i, image in enumerate(images):
                preview_path = file_preview_dir / f"page_{i + 1}.jpg"
                image.save(str(preview_path), "JPEG")
                preview_urls.append(f"/api/previews/{file_path.stem}/{i + 1}")

            return page_count, preview_urls

        except Exception as e:
            print(f"Error processing PDF: {e}")
            return 1, []

    async def process_docx(self, file_path: Path) -> Tuple[int, List[str]]:
        """Process DOCX files."""
        try:
            doc = Document(file_path)
            page_count = len(doc.paragraphs) // 40 + 1  # Rough estimate

            # Create preview directory
            file_preview_dir = self.preview_dir / file_path.stem
            file_preview_dir.mkdir(exist_ok=True)

            # For now, we'll just create a text preview
            preview_path = file_preview_dir / "page_1.txt"
            with open(preview_path, 'w', encoding='utf-8') as f:
                f.write('\n'.join([p.text for p in doc.paragraphs]))

            return page_count, [f"/api/previews/{file_path.stem}/1"]

        except Exception as e:
            print(f"Error processing DOCX: {e}")
            return 1, []

    async def process_text(self, file_path: Path) -> Tuple[int, List[str]]:
        """Process text files."""
        try:
            # Create preview directory
            file_preview_dir = self.preview_dir / file_path.stem
            file_preview_dir.mkdir(exist_ok=True)

            # Copy the text file as preview
            preview_path = file_preview_dir / "page_1.txt"
            shutil.copy2(file_path, preview_path)

            with open(file_path, 'r', encoding='utf-8') as file:
                lines = file.readlines()
                page_count = len(lines) // 50 + 1  # Rough estimate

            return page_count, [f"/api/previews/{file_path.stem}/1"]

        except Exception as e:
            print(f"Error processing text file: {e}")
            return 1, []

    async def get_preview(self, document_id: str, page: int) -> Tuple[str, bytes]:
        """Get preview file content and type."""
        try:
            preview_dir = self.preview_dir / document_id
            
            # Find the preview file
            preview_files = list(preview_dir.glob(f"page_{page}.*"))
            if not preview_files:
                raise FileNotFoundError(f"Preview not found for document {document_id}, page {page}")

            preview_path = preview_files[0]
            
            # Determine content type
            content_type = {
                '.jpg': 'image/jpeg',
                '.png': 'image/png',
                '.txt': 'text/plain'
            }.get(preview_path.suffix, 'application/octet-stream')

            # Read the file content
            with open(preview_path, 'rb') as f:
                content = f.read()

            return content_type, content

        except Exception as e:
            print(f"Error getting preview: {e}")
            raise 